import base64
import io
from datetime import datetime

import pandas as pd
import streamlit as st

from seating_planner import attendees_from_spreadsheet, SeatingOptimizer, SeatingHistory, TableConstraints


def get_download_link(file_path, link_text):
    """Generate a download link for a file"""
    with open(file_path, 'rb') as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{file_path}">{link_text}</a>'


def display_arrangement(arrangement, show_metrics=True):
    """Display the seating arrangement in Streamlit"""
    # Head table
    st.subheader("Head Table")
    head_df = pd.DataFrame([{"Name": a.name, "Gender": a.gender, "Seniority": a.seniority, "Field": a.field,
        "Assigned": "✓" if a.assign_head_table else ""} for a in arrangement[0]])
    st.dataframe(head_df, hide_index=True)

    # Other tables
    for i, table in enumerate(arrangement[1:], 2):
        st.subheader(f"Table {i}")
        table_df = pd.DataFrame(
            [{"Name": a.name, "Gender": a.gender, "Seniority": a.seniority, "Field": a.field} for a in table])
        st.dataframe(table_df, hide_index=True)


def show_excel_guide():
    """Display the guide for Excel file formatting"""
    st.header("Excel File Format Guide")

    st.write("""
    Your Excel file should contain the following columns:
    """)

    # Show example table
    example_data = {'name': ['John Doe', 'Jane Smith', 'Bob Johnson'], 'gender': ['M', 'F', 'M'],
        'seniority': ['senior', 'junior', 'senior'], 'division': ['Sales', 'Engineering', 'Marketing'],
        'attending': ['Y', 'Y', 'Y'], 'head table': [1.0, 0.0, 0.0]}
    example_df = pd.DataFrame(example_data)
    st.dataframe(example_df, hide_index=True)

    st.write("""
    ### Column Descriptions

    1. **name** (required)
        - Full name of the attendee
        - Text format
        - Be sure it matches any existing history data

    2. **seniority** (required)
        - 'senior' or 'junior' or 'guest'
        - Used for mixing experience levels

    3. **division** (required)
        - Department or team name
        - Used for cross-functional mixing
    
    4. **gender** 
        - 'M' or 'F' or 'N'
        - Used for diversity balancing

    5. **attending** 
        - 'Y' or 'N'
        - Only 'Y' entries will be included

    6. **head table**
        - 'TRUE' for assigned head table seats
        - blank for flexible seating

    ### Tips
    - Ensure all columns are present
    - Check for consistent formatting (especially Y/N and 1.0/0.0)
    - Remove any blank rows
    - Save as .xlsx format
    """)

    # Add download link for template
    st.write("### Download Template")

    # Create template DataFrame
    template_df = pd.DataFrame(
        {'name': ['Example Person'], 'seniority': ['senior'], 'division': ['Department'], 'gender': ['F'],
            'attending': ['Y'], 'head table': ['TRUE']})

    # Convert to Excel bytes
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        template_df.to_excel(writer, index=False)

    st.download_button(label="Download Excel Template", data=buffer.getvalue(), file_name="seating_template.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


def generate_word_doc(arrangement):
    """Generate a Word document with the seating arrangement"""
    from docx import Document

    doc = Document()
    doc.add_heading('Seating Arrangement', 0)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

    # Summary table
    doc.add_heading('Summary', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Table'
    header_cells[1].text = 'Number of Guests'

    for i, group in enumerate(arrangement):
        row_cells = table.add_row().cells
        row_cells[0].text = 'Head Table' if i == 0 else f'Table {i + 1}'
        row_cells[1].text = str(len(group))

    doc.add_paragraph()

    # Head table
    doc.add_heading('Head Table', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for i, header in enumerate(['Name', 'Gender', 'Seniority', 'Field', 'Assigned']):
        header_cells[i].text = header

    for attendee in arrangement[0]:
        row_cells = table.add_row().cells
        row_cells[0].text = attendee.name
        row_cells[1].text = attendee.gender
        row_cells[2].text = attendee.seniority
        row_cells[3].text = attendee.field
        row_cells[4].text = '✓' if attendee.assign_head_table else ''

    # Other tables
    for i, table_group in enumerate(arrangement[1:], 2):
        doc.add_paragraph()
        doc.add_heading(f'Table {i}', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        for i, header in enumerate(['Name', 'Gender', 'Seniority', 'Field']):
            header_cells[i].text = header

        for attendee in table_group:
            row_cells = table.add_row().cells
            row_cells[0].text = attendee.name
            row_cells[1].text = attendee.gender
            row_cells[2].text = attendee.seniority
            row_cells[3].text = attendee.field

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def generate_pdf(arrangement):
    """Generate a PDF with the seating arrangement"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet

    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    # Title
    elements.append(Paragraph('Seating Arrangement', styles['Title']))
    elements.append(Paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', styles['Normal']))
    elements.append(Spacer(1, 12))

    # Summary
    elements.append(Paragraph('Summary', styles['Heading1']))
    summary_data = [['Table', 'Number of Guests']]
    for i, group in enumerate(arrangement):
        table_name = 'Head Table' if i == 0 else f'Table {i + 1}'
        summary_data.append([table_name, str(len(group))])

    summary_table = Table(summary_data)
    summary_table.setStyle(TableStyle(
        [('GRID', (0, 0), (-1, -1), 1, colors.black), ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke), ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'), ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12), ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black), ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 12), ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black), ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 20))

    # Head table
    elements.append(Paragraph('Head Table', styles['Heading1']))
    head_data = [['Name', 'Gender', 'Seniority', 'Field', 'Assigned']]
    for attendee in arrangement[0]:
        head_data.append([attendee.name, attendee.gender, attendee.seniority, attendee.field,
            '✓' if attendee.assign_head_table else ''])

    head_table = Table(head_data)
    head_table.setStyle(TableStyle(
        [('GRID', (0, 0), (-1, -1), 1, colors.black), ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke), ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'), ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12), ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black), ('ALIGN', (0, 0), (-1, -1), 'LEFT'), ]))
    elements.append(head_table)
    elements.append(Spacer(1, 20))

    # Other tables
    for i, table_group in enumerate(arrangement[1:], 2):
        elements.append(Paragraph(f'Table {i}', styles['Heading1']))
        table_data = [['Name', 'Gender', 'Seniority', 'Field']]
        for attendee in table_group:
            table_data.append([attendee.name, attendee.gender, attendee.seniority, attendee.field])

        group_table = Table(table_data)
        group_table.setStyle(TableStyle(
            [('GRID', (0, 0), (-1, -1), 1, colors.black), ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke), ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'), ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12), ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black), ('ALIGN', (0, 0), (-1, -1), 'LEFT'), ]))
        elements.append(group_table)
        elements.append(Spacer(1, 20))

    doc.build(elements)
    return output.getvalue()


def main():
    st.set_page_config(page_title="Seating Arrangement Planner", layout="wide")

    # Create tabs
    tab1, tab2 = st.tabs(["Seating Planner", "Attendee Excel File Format Guide"])

    with tab1:
        st.title("Seating Arrangement Planner")

        # Initialize session state
        if 'current_arrangement' not in st.session_state:
            st.session_state.current_arrangement = None
        if 'history' not in st.session_state:
            st.session_state.history = None
        if 'optimizer' not in st.session_state:
            st.session_state.optimizer = None
        if 'previous_history_df' not in st.session_state:
            st.session_state.previous_history_df = None
        if 'previous_min_size' not in st.session_state:
            st.session_state.previous_min_size = None
        if 'previous_max_size' not in st.session_state:
            st.session_state.previous_max_size = None

        # Sidebar for configuration
        with st.sidebar:
            st.header("Configuration")

            # File uploads
            st.write("### Input Files")
            excel_file = st.file_uploader("Upload Attendee Excel File", type=['xlsx'])
            history_file = st.file_uploader("Upload History Excel File (Optional)", type=['xlsx'])

            if history_file:
                try:
                    # Read without setting index
                    st.session_state.previous_history_df = pd.read_excel(history_file)
                    event_columns = [col for col in st.session_state.previous_history_df.columns if col != 'Name']
                    st.success(f"Loaded history with {len(event_columns)} events")
                except Exception as e:
                    st.error(f"Error loading history file: {str(e)}")
            # horizontal rule
            st.markdown("---")

            # Event Date
            event_date = st.date_input("Event Date", datetime.today())

            # Table size configuration
            st.write("### Table Size Configuration")
            min_table_size = st.number_input("Minimum Seats per Table", min_value=4, max_value=8, value=4)
            max_table_size = st.number_input("Maximum Seats per Table", min_value=min_table_size, max_value=12,
                                             value=max(min_table_size, 8))

            # Check if table sizes have changed
            table_sizes_changed = (
                    st.session_state.previous_min_size != min_table_size or st.session_state.previous_max_size != max_table_size)

            if table_sizes_changed:
                # Reset the current arrangement when table sizes change
                st.session_state.current_arrangement = None
                # Update the stored sizes
                st.session_state.previous_min_size = min_table_size
                st.session_state.previous_max_size = max_table_size

            # Calculate and show recommended number of tables
            if excel_file:
                attendees = attendees_from_spreadsheet(excel_file)
                n_guests = len(attendees)
                recommended_tables = max(2, int((n_guests - max_table_size) / max_table_size) + 2)

                st.write("### Table Count")
                st.write(f"Total Guests: {n_guests}")
                st.write(f"Recommended Tables: {recommended_tables}")
                st.write(f"(Based on {min_table_size}-{max_table_size} seats per table)")

                num_tables = st.number_input("Number of Tables (including head table)", min_value=2,
                    value=recommended_tables,
                    help="Default is calculated based on guests/table size. Adjust if needed.")

                # Diversity weights
                st.write("### Diversity Weights")
                gender_weight = st.number_input("Weight for Gender Diversity", value=1.0)
                seniority_weight = st.number_input("Weight for Seniority Diversity", value=1.0)
                field_weight = st.number_input("Weight for Field Diversity", value=1.0)

            # Action buttons
            generate_button = st.button("Generate New Arrangement")

        # Main content area
        if excel_file:
            try:
                # Initialize or update optimizer if needed
                constraints = TableConstraints(min_seats=min_table_size, max_seats=max_table_size)

                if (st.session_state.optimizer is None or table_sizes_changed or history_file):
                    st.session_state.history = SeatingHistory(filename=history_file.name if history_file else None)
                    st.session_state.optimizer = SeatingOptimizer(constraints=constraints,
                        history=st.session_state.history,
                        weights={"gender_weight": gender_weight, "seniority_weight": seniority_weight,
                                 "field_weight": field_weight})

                # Generate new arrangement if requested
                if generate_button or st.session_state.current_arrangement is None:
                    attendees = attendees_from_spreadsheet(excel_file)
                    st.session_state.current_arrangement = st.session_state.optimizer.optimize_seating(attendees,
                        num_tables=num_tables)

                # Display current arrangement
                if st.session_state.current_arrangement:
                    display_arrangement(st.session_state.current_arrangement)

                    # Approval and download section
                    st.write("### Approve and Download")
                    if st.button("Approve & Generate Updated History"):
                        # Create updated history DataFrame
                        new_history_df = st.session_state.history.create_updated_history(
                            st.session_state.current_arrangement,
                            st.session_state.previous_history_df,
                            event_date=event_date
                        )

                        # Convert to Excel for download
                        output = io.BytesIO()
                        with pd.ExcelWriter(output, engine='openpyxl') as writer:
                            new_history_df.to_excel(writer, index=False)  # Don't write index

                        # Offer download
                        st.download_button(
                            "Download Updated History Excel",
                            output.getvalue(),
                            file_name=f"seating_history_{datetime.now().strftime('%Y%m%d')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )

                # Then show export options
                st.write("### Export This Seating Plan")
                col1, col2 = st.columns(2)

                with col1:
                    if st.button("Export to Word"):
                        word_bytes = generate_word_doc(st.session_state.current_arrangement)
                        st.download_button("Download Word Document", word_bytes, file_name="seating_arrangement.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                with col2:
                    if st.button("Export to PDF"):
                        pdf_bytes = generate_pdf(st.session_state.current_arrangement)
                        st.download_button("Download PDF", pdf_bytes, file_name="seating_arrangement.pdf",
                            mime="application/pdf")


            except Exception as e:
                st.error(f"Error: {str(e)}")
                # st.exception(e)
        else:
            st.info("Please upload an Excel file with attendee data to begin.")

    with tab2:
        show_excel_guide()


if __name__ == "__main__":
    main()
